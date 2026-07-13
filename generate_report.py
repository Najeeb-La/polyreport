"""
PolyRapport — Générateur de rapport technique
Polytechnique Montréal · Python pur (python-docx)
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from io import BytesIO
import os

# ─────────────────────────────────────────────
# CONSTANTES
# ─────────────────────────────────────────────

POLY_RED   = RGBColor(0xCC, 0x00, 0x00)
POLY_DARK  = RGBColor(0x1F, 0x1F, 0x1F)
POLY_GRAY  = RGBColor(0x59, 0x59, 0x59)
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)

FONTS_AVAILABLE = [
    "Cambria", "Times New Roman", "Arial", "Calibri",
    "Garamond", "Georgia", "Palatino Linotype", "Book Antiqua",
    "Helvetica", "Verdana", "Tahoma", "Century Schoolbook",
]
REF_STYLES = ["IEEE", "APA", "Chicago"]

MARGIN_CM      = 2.5
HEADER_DIST_CM = 1.25
FOOTER_DIST_CM = 1.25

_DIR      = os.path.dirname(os.path.abspath(__file__))
LOGO_PATH = os.path.join(_DIR, "poly_logo.png")


# ─────────────────────────────────────────────
# HELPERS XML
# ─────────────────────────────────────────────

def _set_margins(section, top=MARGIN_CM, bottom=MARGIN_CM):
    section.top_margin      = Cm(top)
    section.bottom_margin   = Cm(bottom)
    section.left_margin     = Cm(MARGIN_CM)
    section.right_margin    = Cm(MARGIN_CM)
    section.header_distance = Cm(HEADER_DIST_CM)
    section.footer_distance = Cm(FOOTER_DIST_CM)


def _no_borders(table):
    tbl   = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    for e in tblPr.findall(qn('w:tblBorders')): tblPr.remove(e)
    tblBdr = OxmlElement('w:tblBorders')
    for side in ('top','left','bottom','right','insideH','insideV'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'none')
        tblBdr.append(el)
    tblPr.append(tblBdr)


def _shade_cell(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    for e in tcPr.findall(qn('w:shd')): tcPr.remove(e)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  fill_hex)
    tcPr.append(shd)


def _cell_borders(cell, color="AAAAAA", sz=4):
    tcPr  = cell._tc.get_or_add_tcPr()
    for e in tcPr.findall(qn('w:tcBorders')): tcPr.remove(e)
    tcBdr = OxmlElement('w:tcBorders')
    for side in ('top','left','bottom','right'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'),   'single')
        el.set(qn('w:sz'),    str(sz))
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), color)
        tcBdr.append(el)
    tcPr.append(tcBdr)


def _nil_borders(cell):
    tcPr  = cell._tc.get_or_add_tcPr()
    for e in tcPr.findall(qn('w:tcBorders')): tcPr.remove(e)
    tcBdr = OxmlElement('w:tcBorders')
    for side in ('top','left','bottom','right','insideH','insideV'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'nil')
        tcBdr.append(el)
    tcPr.append(tcBdr)


def _run(para, text, font, size_pt, bold=False, italic=False, color=None):
    r = para.add_run(text)
    r.font.name  = font
    r.font.size  = Pt(size_pt)
    r.bold       = bold
    r.italic     = italic
    if color: r.font.color.rgb = color
    return r


def _add_hyperlink(paragraph, url, text, font, size_pt):
    """Vrai lien cliquable Word via relation XML."""
    try:
        part = paragraph.part
        r_id = part.relate_to(
            url,
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
            is_external=True
        )
        hyperlink = OxmlElement('w:hyperlink')
        hyperlink.set(qn('r:id'), r_id)
        new_run = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        color_el = OxmlElement('w:color'); color_el.set(qn('w:val'), '0563C1')
        u_el = OxmlElement('w:u');         u_el.set(qn('w:val'), 'single')
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'), font); rFonts.set(qn('w:hAnsi'), font)
        sz_el = OxmlElement('w:sz'); sz_el.set(qn('w:val'), str(int(size_pt*2)))
        rPr.extend([color_el, u_el, rFonts, sz_el])
        new_run.append(rPr)
        t = OxmlElement('w:t'); t.text = text
        new_run.append(t)
        hyperlink.append(new_run)
        paragraph._p.append(hyperlink)
    except Exception:
        _run(paragraph, text, font, size_pt, color=RGBColor(5,99,193))


def _para(doc, text="", font="Cambria", size=12,
          bold=False, italic=False, color=None,
          align=WD_ALIGN_PARAGRAPH.LEFT,
          sb=0, sa=6, ls=1.5):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before      = Pt(sb)
    p.paragraph_format.space_after       = Pt(sa)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE \
        if ls == 1.5 else WD_LINE_SPACING.MULTIPLE
    if ls != 1.5: p.paragraph_format.line_spacing = ls
    if text: _run(p, text, font, size, bold, italic, color)
    return p


def _page_break(doc):
    p  = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    r  = p.add_run()
    br = OxmlElement('w:br'); br.set(qn('w:type'), 'page')
    r._r.append(br)
    return p


def _field(para, instr, font, size_pt):
    run = para.add_run()
    b   = OxmlElement('w:fldChar'); b.set(qn('w:fldCharType'), 'begin')
    run._r.append(b)
    ins = OxmlElement('w:instrText'); ins.set(qn('xml:space'), 'preserve')
    ins.text = f' {instr} '; run._r.append(ins)
    s = OxmlElement('w:fldChar'); s.set(qn('w:fldCharType'), 'separate')
    run._r.append(s)
    ph = para.add_run("?"); ph.font.name = font; ph.font.size = Pt(size_pt)
    er = para.add_run()
    e  = OxmlElement('w:fldChar'); e.set(qn('w:fldCharType'), 'end')
    er._r.append(e)


def _toc_field(doc, instr, font, placeholder):
    p   = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    run = p.add_run()
    b   = OxmlElement('w:fldChar'); b.set(qn('w:fldCharType'), 'begin')
    b.set(qn('w:dirty'), 'true'); run._r.append(b)
    ins = OxmlElement('w:instrText'); ins.set(qn('xml:space'), 'preserve')
    ins.text = f' {instr} '; run._r.append(ins)
    s = OxmlElement('w:fldChar'); s.set(qn('w:fldCharType'), 'separate')
    run._r.append(s)
    ph = p.add_run(placeholder)
    ph.font.name = font; ph.font.size = Pt(10); ph.italic = True
    ph.font.color.rgb = POLY_GRAY
    er = p.add_run()
    e  = OxmlElement('w:fldChar'); e.set(qn('w:fldCharType'), 'end')
    er._r.append(e)
    return p


def _set_pgnum(section, fmt='decimal', start=1):
    sectPr = section._sectPr
    for e in sectPr.findall(qn('w:pgNumType')): sectPr.remove(e)
    pgNum = OxmlElement('w:pgNumType')
    pgNum.set(qn('w:fmt'),   fmt)
    pgNum.set(qn('w:start'), str(start))
    sectPr.append(pgNum)


def _suppress_pgnum(section):
    """Supprime la numérotation sur une section (page titre)."""
    sectPr = section._sectPr
    for e in sectPr.findall(qn('w:pgNumType')): sectPr.remove(e)


def _section_break(doc, break_type='nextPage', pgnum_fmt=None, pgnum_start=1):
    """Insère un saut de section dans un paragraphe vide."""
    p      = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    pPr    = p._p.get_or_add_pPr()
    sectPr = OxmlElement('w:sectPr')
    typeEl = OxmlElement('w:type'); typeEl.set(qn('w:val'), break_type)
    sectPr.append(typeEl)
    if pgnum_fmt:
        pgNum = OxmlElement('w:pgNumType')
        pgNum.set(qn('w:fmt'),   pgnum_fmt)
        pgNum.set(qn('w:start'), str(pgnum_start))
        sectPr.append(pgNum)
    pPr.append(sectPr)
    return p


# ─────────────────────────────────────────────
# STYLES POLY
# ─────────────────────────────────────────────

def _register_styles(doc, font):
    from docx.enum.style import WD_STYLE_TYPE

    def _get(name):
        try:    return doc.styles[name]
        except: return doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)

    # Body Text
    bt = _get('Body Text')
    bt.font.name = font; bt.font.size = Pt(12)
    bt.paragraph_format.space_before      = Pt(0)
    bt.paragraph_format.space_after       = Pt(6)
    bt.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

    # Références
    cr = _get('CORPS_REFERENCES')
    cr.font.name = font; cr.font.size = Pt(12)
    cr.paragraph_format.space_before      = Pt(0)
    cr.paragraph_format.space_after       = Pt(6)
    cr.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    cr.paragraph_format.left_indent       = Cm(0.75)
    cr.paragraph_format.first_line_indent = Cm(-0.75)

    # Légendes
    lf = _get('LEGENDE_Figure')
    lf.font.name = font; lf.font.size = Pt(10); lf.font.italic = True
    lf.paragraph_format.space_before = Pt(4); lf.paragraph_format.space_after = Pt(10)

    lt = _get('LEGENDE_Tableau')
    lt.font.name = font; lt.font.size = Pt(10); lt.font.bold = True
    lt.paragraph_format.space_before = Pt(10); lt.paragraph_format.space_after = Pt(4)

    # Appliquer les styles Heading 1–4 nativement Word pour que la TDM fonctionne
    _style_heading(doc, 'Heading 1', font, 14, bold=True,  sb=18, sa=6)
    _style_heading(doc, 'Heading 2', font, 12, bold=True,  sb=12, sa=6)
    _style_heading(doc, 'Heading 3', font, 12, bold=True,  italic=True, sb=10, sa=6)
    _style_heading(doc, 'Heading 4', font, 12, italic=True, sb=10, sa=6)


def _style_heading(doc, style_name, font, size, bold=False, italic=False, sb=0, sa=6):
    try:
        s = doc.styles[style_name]
        s.font.name   = font
        s.font.size   = Pt(size)
        s.font.bold   = bold
        s.font.italic = italic
        s.font.color.rgb = POLY_DARK
        s.paragraph_format.space_before      = Pt(sb)
        s.paragraph_format.space_after       = Pt(sa)
        s.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    except Exception:
        pass


# ─────────────────────────────────────────────
# EN-TÊTE — simple et professionnel
# ─────────────────────────────────────────────

def _build_header(section, config, font="Cambria", show=True):
    header = section.header
    header.is_linked_to_previous = False
    for p in header.paragraphs: p.clear()

    if not show:
        return  # page titre : en-tête vide

    report_title = config.get("report_title", "").strip()

    tbl = header.add_table(1, 2, width=Inches(6.3))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    _no_borders(tbl)
    tbl.cell(0,0).width = Inches(4.8)
    tbl.cell(0,1).width = Inches(1.5)
    _nil_borders(tbl.cell(0,0)); _nil_borders(tbl.cell(0,1))

    fs = 9
    p_left = tbl.cell(0,0).paragraphs[0]
    r_left = p_left.add_run(report_title[:70] if report_title else "")
    r_left.font.name = font; r_left.font.size = Pt(fs)
    r_left.font.color.rgb = POLY_GRAY

    p_right = tbl.cell(0,1).paragraphs[0]
    p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r_pre = p_right.add_run("Page ")
    r_pre.font.name = font; r_pre.font.size = Pt(fs); r_pre.font.color.rgb = POLY_GRAY
    _field(p_right, 'PAGE \\* MERGEFORMAT', font, fs)
    r_mid = p_right.add_run(" / ")
    r_mid.font.name = font; r_mid.font.size = Pt(fs); r_mid.font.color.rgb = POLY_GRAY
    _field(p_right, 'SECTIONPAGES \\* MERGEFORMAT', font, fs)

    p_rule = header.add_paragraph()
    p_rule.paragraph_format.space_before = Pt(3)
    p_rule.paragraph_format.space_after  = Pt(0)
    pPr  = p_rule._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot  = OxmlElement('w:bottom')
    bot.set(qn('w:val'),   'single'); bot.set(qn('w:sz'), '4')
    bot.set(qn('w:space'), '1');      bot.set(qn('w:color'), 'CCCCCC')
    pBdr.append(bot); pPr.append(pBdr)


# ─────────────────────────────────────────────
# PAGE TITRE — pleine page, sans numéro, logo centré
# ─────────────────────────────────────────────

def build_title_page(doc, config):
    font         = config.get("font", "Cambria")
    course_code  = config.get("course_code", "").strip()
    course_name  = config.get("course_name", "").strip()
    group_number = config.get("group_number", "").strip()
    report_title = config.get("report_title", "").strip()
    professors   = [p.strip() for p in config.get("professors", []) if str(p).strip()]
    members      = [m for m in config.get("members", [])
                    if (m.get("name","") if isinstance(m,dict) else str(m)).strip()]
    doc_date     = config.get("doc_date", "").strip()
    is_individual= config.get("is_individual", False)
    semester     = config.get("semester", "").strip()

    section = doc.sections[0]
    _set_margins(section, top=2.5, bottom=2.5)
    _build_header(section, config, font, show=False)  # en-tête vide sur page titre
    _suppress_pgnum(section)  # pas de numéro sur la page titre

    # ── Logo centré ──────────────────────────
    sp = Pt(0)
    if os.path.exists(LOGO_PATH):
        p_logo = doc.add_paragraph()
        p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_logo.paragraph_format.space_before = Pt(40)
        p_logo.paragraph_format.space_after  = Pt(30)
        run_logo = p_logo.add_run()
        run_logo.add_picture(LOGO_PATH, width=Inches(2.6))
    else:
        sp = Pt(60)

    # ── Cours ────────────────────────────────
    if course_code or course_name:
        txt = " — ".join(filter(None, [course_code, course_name]))
        p_c = doc.add_paragraph()
        p_c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_c.paragraph_format.space_before = sp
        p_c.paragraph_format.space_after  = Pt(2)
        _run(p_c, txt, font, 12, bold=True)

    if group_number:
        p_g = doc.add_paragraph()
        p_g.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_g.paragraph_format.space_before = Pt(0)
        p_g.paragraph_format.space_after  = Pt(2)
        _run(p_g, f"Groupe-cours n° {group_number}", font, 12)

    if semester:
        p_s = doc.add_paragraph()
        p_s.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_s.paragraph_format.space_before = Pt(0)
        p_s.paragraph_format.space_after  = Pt(30)
        _run(p_s, semester, font, 12)
    else:
        p_sp = doc.add_paragraph()
        p_sp.paragraph_format.space_after = Pt(30)

    # ── Titre du rapport ─────────────────────
    if report_title:
        p_t = doc.add_paragraph()
        p_t.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_t.paragraph_format.space_before = Pt(0)
        p_t.paragraph_format.space_after  = Pt(40)
        _run(p_t, report_title, font, 14)

    # ── Présenté à — une ligne vide puis un prof par ligne ──
    if professors:
        p_pr_label = doc.add_paragraph()
        p_pr_label.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_pr_label.paragraph_format.space_before = Pt(0)
        p_pr_label.paragraph_format.space_after  = Pt(6)
        _run(p_pr_label, "Présenté à", font, 12)

        # Ligne vide
        p_blank = doc.add_paragraph()
        p_blank.paragraph_format.space_before = Pt(0)
        p_blank.paragraph_format.space_after  = Pt(0)

        for prof in professors:
            p_p = doc.add_paragraph()
            p_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_p.paragraph_format.space_before = Pt(0)
            p_p.paragraph_format.space_after  = Pt(4)
            _run(p_p, prof, font, 12)

    # ── Fait par ─────────────────────────────
    if is_individual and members:
        p_fb = doc.add_paragraph()
        p_fb.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_fb.paragraph_format.space_before = Pt(30)
        p_fb.paragraph_format.space_after  = Pt(6)
        _run(p_fb, "Fait par", font, 12)
        m   = members[0]
        nom = m.get("name","") if isinstance(m,dict) else str(m)
        mat = m.get("matricule","") if isinstance(m,dict) else ""
        p_m = doc.add_paragraph()
        p_m.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_m.paragraph_format.space_before = Pt(0)
        p_m.paragraph_format.space_after  = Pt(4)
        _run(p_m, f"{nom}  {mat}".strip(), font, 12)

    elif not is_individual:
        p_fb = doc.add_paragraph()
        p_fb.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_fb.paragraph_format.space_before = Pt(30)
        p_fb.paragraph_format.space_after  = Pt(12)
        _run(p_fb, "Fait par", font, 12)

        # Tableau avec signatures (inchangé)
        n       = max(len(members), 4)
        sig_tbl = doc.add_table(rows=n+1, cols=3)
        sig_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr_w   = [Inches(2.0), Inches(1.0), Inches(3.0)]
        hdrs    = ["Nom, Prénom", "Matricule", "_______Signatures requises_______"]
        for j, (h, w) in enumerate(zip(hdrs, hdr_w)):
            c = sig_tbl.cell(0, j); c.width = w
            _shade_cell(c, "1F1F1F"); _cell_borders(c, "1F1F1F")
            p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(h); r.font.name=font; r.font.size=Pt(10)
            r.bold=True; r.font.color.rgb=WHITE
        for i in range(n):
            row = sig_tbl.rows[i+1]
            m   = members[i] if i < len(members) else None
            nom = (m.get("name","") if isinstance(m,dict) else str(m)) if m else ""
            mat = (m.get("matricule","") if isinstance(m,dict) else "") if m else ""
            for j, (val, w) in enumerate(zip([nom, mat, "_"*42], hdr_w)):
                c = row.cells[j]; c.width=w; _cell_borders(c,"AAAAAA")
                p = c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.LEFT
                _run(p, val, font, 10)

    # ── Date ─────────────────────────────────
    if doc_date:
        p_d = doc.add_paragraph()
        p_d.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_d.paragraph_format.space_before = Pt(30)
        p_d.paragraph_format.space_after  = Pt(0)
        _run(p_d, f"Le {doc_date}", font, 12)
        p_m2 = doc.add_paragraph()
        p_m2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_m2.paragraph_format.space_before = Pt(0)
        p_m2.paragraph_format.space_after  = Pt(0)
        _run(p_m2, "à Montréal", font, 12)


# ─────────────────────────────────────────────
# TITRE DE PAGE PRÉLIMINAIRE — centré, haut de page
# ─────────────────────────────────────────────

def _prelim_title(doc, title, font):
    """Titre centré en haut pour TDM, listes, références, annexes."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(18)
    r = p.add_run(title)
    r.font.name = font; r.font.size = Pt(16); r.bold = True


# ─────────────────────────────────────────────
# RÉSUMÉ + ABSTRACT — champs optionnels
# ─────────────────────────────────────────────

def build_resume_abstract(doc, config):
    font        = config.get("font", "Cambria")
    size        = config.get("font_size", 12)
    resume_fr   = config.get("resume_fr", "").strip()
    abstract_en = config.get("abstract_en", "").strip()
    keywords_fr = config.get("keywords_fr", "").strip()
    keywords_en = config.get("keywords_en", "").strip()

    PLACEHOLDER = (
        "[ À rédiger après la complétion du rapport. "
        "Supprimer ce texte et le remplacer par votre résumé. ]"
    )
    PLACEHOLDER_EN = (
        "[ To be written after the report is completed. "
        "Delete this text and replace with your abstract. ]"
    )

    def sec(title):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after  = Pt(6)
        _run(p, title, font, size, bold=True)

    sec("Résumé")
    try:   p = doc.add_paragraph(style='Body Text')
    except: p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    txt = resume_fr if resume_fr else PLACEHOLDER
    r   = _run(p, txt, font, size)
    if not resume_fr: r.italic = True; r.font.color.rgb = POLY_GRAY

    if keywords_fr:
        try:   pk = doc.add_paragraph(style='Body Text')
        except: pk = doc.add_paragraph()
        pk.paragraph_format.space_after = Pt(12)
        rb = pk.add_run("Mots clés : "); rb.font.name=font; rb.font.size=Pt(size); rb.bold=True
        pk.add_run(keywords_fr).font.name = font

    sec("Abstract")
    try:   p = doc.add_paragraph(style='Body Text')
    except: p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    txt = abstract_en if abstract_en else PLACEHOLDER_EN
    r   = _run(p, txt, font, size, italic=(not abstract_en))
    if not abstract_en: r.font.color.rgb = POLY_GRAY

    if keywords_en:
        try:   pk = doc.add_paragraph(style='Body Text')
        except: pk = doc.add_paragraph()
        pk.paragraph_format.space_after = Pt(0)
        rb = pk.add_run("Keywords: "); rb.font.name=font; rb.font.size=Pt(size); rb.bold=True
        pk.add_run(keywords_en).font.name = font

    _page_break(doc)


# ─────────────────────────────────────────────
# TABLE DES MATIÈRES + LISTE FIGURES + LISTE TABLEAUX
# ─────────────────────────────────────────────

def build_toc_section(doc, config):
    font = config.get("font", "Cambria")

    _prelim_title(doc, "Table des matières", font)
    _toc_field(doc, 'TOC \\o "1-4" \\h \\z \\u', font,
               "[ Ctrl+A → F9 dans Word pour mettre à jour ]")
    _page_break(doc)

    _prelim_title(doc, "Liste des tableaux", font)
    _toc_field(doc, 'TOC \\h \\z \\c "Tableau"', font,
               "[ Liste des tableaux — F9 pour mettre à jour ]")
    _page_break(doc)

    _prelim_title(doc, "Liste des figures", font)
    _toc_field(doc, 'TOC \\h \\z \\c "Figure"', font,
               "[ Liste des figures — F9 pour mettre à jour ]")
    _page_break(doc)


# ─────────────────────────────────────────────
# TITRES DE SECTIONS — styles Word natifs Heading 1/2/3/4
# ─────────────────────────────────────────────

def _heading(doc, text, level, font, number=""):
    """
    Utilise les styles Word natifs Heading 1/2/3/4 pour que
    la table des matières se mette à jour automatiquement.
    """
    style_name = f"Heading {level}"
    try:
        p = doc.add_paragraph(style=style_name)
    except Exception:
        p = doc.add_paragraph()

    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(18 if level == 1 else 12)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

    full = f"{number}\t{text}" if number else text
    r = p.add_run(full)
    r.font.name  = font
    r.font.size  = Pt(14 if level == 1 else 12)
    r.bold       = level <= 2
    r.italic     = level == 3
    r.font.color.rgb = POLY_DARK
    return p


def _placeholder_text(doc, font, size):
    try:   p = doc.add_paragraph(style='Body Text')
    except: p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(
        "[ Rédiger le contenu de cette section directement dans Word. "
        "Supprimer ce texte gris avant de remettre le rapport. ]"
    )
    r.font.name=font; r.font.size=Pt(size)
    r.italic=True; r.font.color.rgb=POLY_GRAY
    return p


def _insert_figure_placeholder(doc, font, size, fig_num, caption=""):
    tbl = doc.add_table(1, 1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    c = tbl.cell(0, 0); c.width = Inches(5.0)
    _shade_cell(c, "F2F2F2"); _cell_borders(c, "CCCCCC", sz=4)
    cp = c.paragraphs[0]
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp.paragraph_format.space_before = Pt(24)
    cp.paragraph_format.space_after  = Pt(24)
    r = cp.add_run(f"[ Insérer la Figure {fig_num} ici ]")
    r.font.name=font; r.font.size=Pt(10); r.italic=True; r.font.color.rgb=POLY_GRAY

    cap_text = caption if caption else f"Description de la figure {fig_num}."
    try:   p_cap = doc.add_paragraph(style='LEGENDE_Figure')
    except: p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cap.paragraph_format.space_before = Pt(4)
    p_cap.paragraph_format.space_after  = Pt(12)

    r_pre = p_cap.add_run("Figure "); r_pre.font.name=font; r_pre.font.size=Pt(10); r_pre.italic=True
    seq_run = p_cap.add_run()
    b = OxmlElement('w:fldChar'); b.set(qn('w:fldCharType'),'begin'); seq_run._r.append(b)
    ins = OxmlElement('w:instrText'); ins.set(qn('xml:space'),'preserve')
    ins.text = ' SEQ Figure \\* ARABIC '; seq_run._r.append(ins)
    s = OxmlElement('w:fldChar'); s.set(qn('w:fldCharType'),'separate'); seq_run._r.append(s)
    ph = p_cap.add_run(str(fig_num)); ph.font.name=font; ph.font.size=Pt(10); ph.italic=True
    er = p_cap.add_run()
    e = OxmlElement('w:fldChar'); e.set(qn('w:fldCharType'),'end'); er._r.append(e)
    r_cap = p_cap.add_run(f" — {cap_text}")
    r_cap.font.name=font; r_cap.font.size=Pt(10); r_cap.italic=True


def _insert_table_placeholder(doc, font, size, tbl_num, caption="", rows=3, cols=3):
    """Tableau placeholder — nombre de lignes/colonnes configurable."""
    cap_text = caption if caption else f"Description du tableau {tbl_num}."
    try:   p_cap = doc.add_paragraph(style='LEGENDE_Tableau')
    except: p_cap = doc.add_paragraph()
    p_cap.paragraph_format.space_before = Pt(10)
    p_cap.paragraph_format.space_after  = Pt(4)

    r_pre = p_cap.add_run("Tableau "); r_pre.font.name=font; r_pre.font.size=Pt(10); r_pre.bold=True
    seq_run = p_cap.add_run()
    b = OxmlElement('w:fldChar'); b.set(qn('w:fldCharType'),'begin'); seq_run._r.append(b)
    ins = OxmlElement('w:instrText'); ins.set(qn('xml:space'),'preserve')
    ins.text = ' SEQ Tableau \\* ARABIC '; seq_run._r.append(ins)
    s = OxmlElement('w:fldChar'); s.set(qn('w:fldCharType'),'separate'); seq_run._r.append(s)
    ph = p_cap.add_run(str(tbl_num)); ph.font.name=font; ph.font.size=Pt(10); ph.bold=True
    er = p_cap.add_run()
    e = OxmlElement('w:fldChar'); e.set(qn('w:fldCharType'),'end'); er._r.append(e)
    r_cap = p_cap.add_run(f" : {cap_text}")
    r_cap.font.name=font; r_cap.font.size=Pt(10); r_cap.bold=True

    # Tableau avec le bon nombre de lignes/colonnes
    data_tbl = doc.add_table(rows=rows, cols=cols)
    data_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j in range(cols):
        c = data_tbl.cell(0, j)
        _shade_cell(c, "1F1F1F"); _cell_borders(c, "1F1F1F")
        p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(f"En-tête {j+1}"); r.font.name=font; r.font.size=Pt(10)
        r.bold=True; r.font.color.rgb=WHITE
    for i in range(1, rows):
        for j in range(cols):
            c = data_tbl.cell(i,j); _cell_borders(c,"AAAAAA")
            if i % 2 == 0: _shade_cell(c, "F7F7F7")
            _run(c.paragraphs[0], "—", font, 10)

    doc.add_paragraph().paragraph_format.space_after = Pt(10)


# ─────────────────────────────────────────────
# CORPS DU RAPPORT
# ─────────────────────────────────────────────

def build_body_sections(doc, config):
    font  = config.get("font", "Cambria")
    size  = config.get("font_size", 12)
    secs  = config.get("sections", [])

    if not secs:
        secs = [
            {"title":"Introduction",         "level":1,"figures":0,"tables":0,"tbl_rows":3,"tbl_cols":3,"fig_captions":[],"tbl_captions":[]},
            {"title":"Mise en contexte",     "level":2,"figures":0,"tables":0,"tbl_rows":3,"tbl_cols":3,"fig_captions":[],"tbl_captions":[]},
            {"title":"Objectifs",            "level":2,"figures":0,"tables":0,"tbl_rows":3,"tbl_cols":3,"fig_captions":[],"tbl_captions":[]},
            {"title":"Développement",        "level":1,"figures":0,"tables":0,"tbl_rows":3,"tbl_cols":3,"fig_captions":[],"tbl_captions":[]},
            {"title":"Résultats",            "level":1,"figures":0,"tables":0,"tbl_rows":3,"tbl_cols":3,"fig_captions":[],"tbl_captions":[]},
            {"title":"Conclusion",           "level":1,"figures":0,"tables":0,"tbl_rows":3,"tbl_cols":3,"fig_captions":[],"tbl_captions":[]},
        ]

    counters  = [0,0,0,0]
    fig_count = 0
    tbl_count = 0

    for sec in secs:
        lvl   = sec.get("level",1) - 1
        title = sec.get("title","Section sans titre")
        n_fig = int(sec.get("figures", 0))
        n_tbl = int(sec.get("tables",  0))
        t_rows = int(sec.get("tbl_rows", 3))
        t_cols = int(sec.get("tbl_cols", 3))
        fig_caps = sec.get("fig_captions", []) or []
        tbl_caps = sec.get("tbl_captions", []) or []

        counters[lvl] += 1
        for i in range(lvl+1, 4): counters[i] = 0
        num = ".".join(str(counters[i]) for i in range(lvl+1))

        _heading(doc, title, lvl+1, font, number=num)
        _placeholder_text(doc, font, size)

        for fi in range(n_fig):
            fig_count += 1
            cap = fig_caps[fi] if fi < len(fig_caps) else ""
            _insert_figure_placeholder(doc, font, size, fig_count, cap)

        for ti in range(n_tbl):
            tbl_count += 1
            cap = tbl_caps[ti] if ti < len(tbl_caps) else ""
            _insert_table_placeholder(doc, font, size, tbl_count, cap, t_rows, t_cols)

    return fig_count, tbl_count


# ─────────────────────────────────────────────
# RÉFÉRENCES — avec vrai lien cliquable
# ─────────────────────────────────────────────

def build_references(doc, config):
    font       = config.get("font", "Cambria")
    size       = config.get("font_size", 12)
    ref_style  = config.get("ref_style", "IEEE")
    references = config.get("references", [])

    _prelim_title(doc, "Références", font)

    if not references:
        _placeholder_text(doc, font, size)
        return

    for i, ref in enumerate(references, 1):
        formatted = ref.get("formatted","") if isinstance(ref,dict) else str(ref)
        url       = ref.get("url","")       if isinstance(ref,dict) else ""
        if not formatted.strip(): continue

        try:   p = doc.add_paragraph(style='CORPS_REFERENCES')
        except:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent      = Cm(0.75)
            p.paragraph_format.first_line_indent = Cm(-0.75)

        prefix = f"[{i}]\t" if ref_style=="IEEE" else ""
        _run(p, prefix + formatted, font, size)

        if url:
            sep = p.add_run("  —  ")
            sep.font.name=font; sep.font.size=Pt(size-1); sep.font.color.rgb=POLY_GRAY
            _add_hyperlink(p, url, url, font, size-1)


# ─────────────────────────────────────────────
# ANNEXES — numérotation de pages indépendante
# ─────────────────────────────────────────────

def build_annexes(doc, config, fig_count_start=0, tbl_count_start=0):
    font    = config.get("font", "Cambria")
    size    = config.get("font_size", 12)
    annexes = config.get("annexes", [])
    if not annexes: return

    letters   = list("ABCDEFGHIJKLMNOPQRSTUVWXYZ")
    fig_count = fig_count_start
    tbl_count = tbl_count_start

    for i, annexe in enumerate(annexes):
        letter  = letters[i] if i < len(letters) else str(i+1)
        title   = annexe.get("title", f"Titre de l'annexe {letter}")
        n_fig   = int(annexe.get("figures", 0))
        n_tbl   = int(annexe.get("tables",  0))
        t_rows  = int(annexe.get("tbl_rows", 3))
        t_cols  = int(annexe.get("tbl_cols", 3))
        fig_caps= annexe.get("fig_captions", []) or []
        tbl_caps= annexe.get("tbl_captions", []) or []

        # Saut de section : numérotation repart à 1 pour chaque annexe
        _section_break(doc, 'nextPage', pgnum_fmt='decimal', pgnum_start=1)

        _prelim_title(doc, f"Annexe {letter} : {title}", font)
        _placeholder_text(doc, font, size)

        for fi in range(n_fig):
            fig_count += 1
            cap = fig_caps[fi] if fi < len(fig_caps) else ""
            _insert_figure_placeholder(doc, font, size, fig_count, cap)

        for ti in range(n_tbl):
            tbl_count += 1
            cap = tbl_caps[ti] if ti < len(tbl_caps) else ""
            _insert_table_placeholder(doc, font, size, tbl_count, cap, t_rows, t_cols)


# ─────────────────────────────────────────────
# FONCTION PRINCIPALE
# ─────────────────────────────────────────────

def generate_poly_report(config: dict) -> BytesIO:
    """
    Génère un rapport technique Polytechnique Montréal.

    Numérotation des pages :
      - Page titre      : aucun numéro
      - Pages préliminaires (résumé, TDM…) : chiffres romains (i, ii, iii…)
      - Corps du rapport à partir de l'Introduction : arabes (1, 2, 3…)
      - Annexes : arabes indépendants par annexe (1, 2…)
    """
    doc       = Document()
    font_name = config.get("font", "Cambria")
    doc.styles['Normal'].font.name = font_name
    doc.styles['Normal'].font.size = Pt(config.get("font_size", 12))
    _register_styles(doc, font_name)

    # ── 1. Page titre (Section 0 — aucun numéro) ──
    build_title_page(doc, config)

    # ── 2. Section préliminaire — chiffres romains ──
    _section_break(doc, 'nextPage', pgnum_fmt='lowerRoman', pgnum_start=1)
    sec_prelim = doc.sections[-1]
    _set_margins(sec_prelim)
    _build_header(sec_prelim, config, font_name, show=True)

    build_resume_abstract(doc, config)
    build_toc_section(doc, config)

    # ── 3. Corps — chiffres arabes à partir de 1 ──
    _section_break(doc, 'nextPage', pgnum_fmt='decimal', pgnum_start=1)
    sec_body = doc.sections[-1]
    _set_margins(sec_body)
    _build_header(sec_body, config, font_name, show=True)

    fig_count, tbl_count = build_body_sections(doc, config)
    build_references(doc, config)

    # ── 4. Annexes — numérotation indépendante ──
    build_annexes(doc, config, fig_count, tbl_count)

    # Marges sur toutes les sections restantes
    for s in doc.sections:
        if s.left_margin is None or s.left_margin.cm < 1:
            _set_margins(s)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


if __name__ == "__main__":
    cfg = {
        "font":"Cambria","font_size":12,"ref_style":"IEEE",
        "doc_date":"17 juin 2026","course_code":"MEC 3520",
        "course_name":"Industrialisation des produits","group_number":"02",
        "report_title":"Test de validation PolyRapport","semester":"Hiver 2026",
        "professors":["Prof. Marwan Azzi","Dr. Leila Ahmadi"],
        "members":[{"name":"Étudiant Test","matricule":"2000000"}],
        "is_individual":False,
        "resume_fr":"","abstract_en":"",
        "keywords_fr":"","keywords_en":"",
        "sections":[
            {"title":"Introduction","level":1,"figures":1,"tables":1,
             "tbl_rows":4,"tbl_cols":3,
             "fig_captions":["Schéma général."],"tbl_captions":["Données mesurées."]},
            {"title":"Mise en contexte","level":2,"figures":0,"tables":0,
             "tbl_rows":3,"tbl_cols":3,"fig_captions":[],"tbl_captions":[]},
            {"title":"Conclusion","level":1,"figures":0,"tables":0,
             "tbl_rows":3,"tbl_cols":3,"fig_captions":[],"tbl_captions":[]},
        ],
        "references":[
            {"formatted":"Auteur, A. (2024). Titre. Journal, 10(2), 123.",
             "url":"https://doi.org/10.1234/test"},
            {"formatted":"Polytechnique Montréal. (2025). Planeur de secours.",
             "url":"https://moodle.polymtl.ca/pluginfile.php/1449908/mod_resource/content/1/AER2110.pdf"},
        ],
        "annexes":[
            {"title":"Données brutes","figures":0,"tables":1,
             "tbl_rows":5,"tbl_cols":4,
             "fig_captions":[],"tbl_captions":["Mesures complètes."]},
        ],
    }
    buf = generate_poly_report(cfg)
    with open("/home/claude/rapport_test.docx","wb") as f: f.write(buf.read())
    print("OK rapport_test.docx")
